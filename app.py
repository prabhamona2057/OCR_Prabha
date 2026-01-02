import streamlit as st
import pytesseract
from PIL import Image
from docx import Document
import io
import re

# Function to remove characters that break XML/DOCX files
def sanitize_text(text):
    if not text:
        return ""
    # Remove control characters except for newline and tab
    return re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\xff]', '', text)

st.set_page_config(page_title="Image to DOCX Converter", layout="centered")

st.title("📄 Image to Word Converter")
st.write("Upload an image to extract text and download it as a .docx file.")

uploaded_file = st.file_uploader("Choose an image...", type=["png", "jpg", "jpeg"])

if uploaded_file is not None:
    image = Image.open(uploaded_file)
    st.image(image, caption='Uploaded Image', use_container_width=True)
    
    with st.spinner('Extracting and cleaning text...'):
        # 1. Get raw text from OCR
        raw_text = pytesseract.image_to_string(image)
        
        # 2. Sanitize text to fix the ValueError
        clean_text = sanitize_text(raw_text)
        
    if clean_text.strip():
        st.subheader("Extracted Text:")
        st.text_area("Result", clean_text, height=250)
        
        # Create Word Document
        doc = Document()
        doc.add_heading('Extracted Text', 0)
        
        # Add the cleaned text
        doc.add_paragraph(clean_text)
        
        # Save to buffer
        bio = io.BytesIO()
        doc.save(bio)
        
        st.download_button(
            label="Download as DOCX",
            data=bio.getvalue(),
            file_name="extracted_text.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    else:
        st.warning("No readable text detected in the image.")
